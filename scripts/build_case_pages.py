from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "网站素材补充" / "线下融合 专区详细页面"
OUTPUT_DIR = ROOT / "site" / "pages" / "cases"


CASES = [
    {
        "prefix": "案例一",
        "slug": "yuhua-jinhua",
        "label": "案例一 · 南京雨花近华",
        "audience": "社区儿童与家长",
        "summary": "戏曲展演、行当科普、昆曲水袖、脸谱手绘与皮影体验",
        "subheads": {1},
    },
    {
        "prefix": "案例二",
        "slug": "yadong-community",
        "label": "案例二 · 南京亚东社区",
        "audience": "社区居民与银发戏迷",
        "summary": "京剧教唱、经典诗词合诵与跨代邻里文化交流",
        "subheads": {2, 5, 8, 11},
    },
    {
        "prefix": "案例三",
        "slug": "yongmei-community",
        "label": "案例三 · 南京咏梅社区",
        "audience": "中老年戏曲爱好者",
        "summary": "戏曲 AI 宣讲、适老功能演示、语音点戏与需求交流",
        "subheads": set(),
    },
    {
        "prefix": "案例四",
        "slug": "tanzhuang-village",
        "label": "案例四 · 盐城潭庄村",
        "audience": "乡村少年儿童",
        "summary": "戏曲行当启蒙、昆曲身段、脸谱手绘与数字平台延伸",
        "subheads": set(),
    },
]


def photo_path(part_name: str, slug: str) -> str:
    match = re.search(r"image(\d+)", part_name)
    if not match:
        raise ValueError(part_name)
    index = int(match.group(1))
    suffix = Path(part_name).suffix.lower() or ".jpg"
    return f"../../assets/images/cases/{slug}/photo-{index:02d}{suffix}"


def article_body(doc: Document, slug: str, subheads: set[int]) -> str:
    blocks: list[str] = []
    for index, paragraph in enumerate(doc.paragraphs[1:], start=1):
        text = paragraph.text.strip()
        embeds = paragraph._p.xpath(".//a:blip/@r:embed")
        photos = []
        for rel_id in embeds:
            part = doc.part.related_parts.get(rel_id)
            if part:
                photos.append(photo_path(str(part.partname), slug))

        escaped = html.escape(text)
        is_caption = bool(re.match(r"^图\s*\d+", text))

        if index in subheads and text:
            blocks.append(f"<h2>{escaped}</h2>")
        elif text and not is_caption:
            extra_class = ' class="case-byline"' if text.startswith("作者：") else ""
            blocks.append(f"<p{extra_class}>{escaped}</p>")

        if photos:
            for photo_index, source in enumerate(photos):
                caption = escaped if is_caption and photo_index == 0 else ""
                figcaption = f"<figcaption>{caption}</figcaption>" if caption else ""
                alt = caption or f"{html.escape(doc.paragraphs[0].text)}活动照片"
                blocks.append(
                    f'<figure class="case-figure"><img src="{source}" alt="{alt}" loading="lazy">{figcaption}</figure>'
                )
        elif is_caption:
            blocks.append(f'<p class="case-caption">{escaped}</p>')

    return "\n".join(blocks)


def nav() -> str:
    return """
<nav class="topnav" id="topnav">
  <div class="topnav-inner">
    <a class="brand" href="../../index.html"><span class="brand-mark">戏</span><span class="brand-text">数智·戏曲润心</span></a>
    <ul class="nav-links">
      <li><a href="../../index.html">首页</a></li><li><a href="../ai-opera.html">共创工坊</a></li>
      <li><a href="../characters.html">数智人物</a></li><li><a href="../kids.html">儿童专区</a></li>
      <li><a href="../elderly.html">老人陪伴</a></li><li><a href="../red-opera.html">红色戏曲</a></li>
      <li><a href="../exchange.html">文明互鉴</a></li><li><a class="active" href="../live.html">线下融合</a></li>
    </ul>
    <button class="nav-toggle" id="navToggle" aria-label="展开菜单"><span></span><span></span><span></span></button>
  </div>
</nav>"""


def footer() -> str:
    return """
<footer class="footer"><div class="container"><div class="footer-grid">
  <div><div class="footer-brand"><span class="brand-mark">戏</span><span>数智·戏曲润心</span></div><p class="footer-desc">以中华优秀传统戏曲为内核，以数字工具拓展基层文化服务的参与方式。</p><div class="footer-political"><span class="ppill">传统文化活化</span><span class="ppill">数智赋能</span><span class="ppill">基层实践</span><span class="ppill">文明互鉴</span></div></div>
  <div><h5>平台入口</h5><ul><li><a href="../ai-opera.html">戏曲共创工坊</a></li><li><a href="../kids.html">儿童戏曲知识专区</a></li><li><a href="../elderly.html">老人陪伴式戏曲演员</a></li><li><a href="../red-opera.html">红色戏曲传唱</a></li></ul></div>
  <div><h5>项目档案</h5><ul><li>类型：暑期社会实践</li><li>阵地：新时代文明实践站</li><li>形态：线上体验 + 线下微课堂</li><li><a href="../notice.html">内容与使用说明</a></li></ul></div>
</div><div class="footer-bottom"><p>© 2026 数智·戏曲润心 项目组</p><p class="footer-tech">案例文字与活动照片来自项目实践记录。</p></div></div></footer>"""


def build_page(case: dict) -> str:
    matches = list(SOURCE_DIR.glob(f"{case['prefix']}*.docx"))
    if len(matches) != 1:
        raise RuntimeError(f"Expected one source for {case['prefix']}, found {matches}")
    doc = Document(matches[0])
    title = doc.paragraphs[0].text.strip()
    body = article_body(doc, case["slug"], case["subheads"])
    return f"""<!DOCTYPE html>
<html lang="zh-CN"><head>
  <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
  <meta name="theme-color" content="#b31219"><link rel="icon" href="../../assets/icons/mark.svg" type="image/svg+xml">
  <title>{html.escape(title)} | 数智·戏曲润心</title>
  <meta name="description" content="{html.escape(case['summary'])}">
  <link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=ZCOOL+XiaoWei&family=Ma+Shan+Zheng&family=Noto+Serif+SC:wght@400;700&family=Noto+Sans+SC:wght@300;400;500;700&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="../../css/style.css">
</head><body>
{nav()}
<header class="case-article-hero">
  <img src="../../assets/images/cases/{case['slug']}-cover.jpg" alt="{html.escape(case['label'])}活动封面">
  <div class="case-article-shade"></div>
  <div class="container case-article-hero-inner"><div class="hero-badge">{html.escape(case['label'])}</div><h1>{html.escape(title)}</h1><p>{html.escape(case['audience'])} · {html.escape(case['summary'])}</p></div>
</header>
<main class="section case-article-section"><article class="container case-article">
  <a class="case-back" href="../live.html">← 返回四站实践纪实</a>
  {body}
  <div class="back-row"><a class="btn btn-ghost" href="../live.html">← 返回案例总览</a></div>
</article></main>
{footer()}
<script src="../../js/shared.js"></script>
</body></html>"""


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for case in CASES:
        output = OUTPUT_DIR / f"{case['slug']}.html"
        output.write_text(build_page(case), encoding="utf-8", newline="\n")
        print(output.relative_to(ROOT))


if __name__ == "__main__":
    main()
